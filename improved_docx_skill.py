#!/usr/bin/env python3
"""
Improved Docx Skill - Read and write Microsoft Word (.docx) files
Uses manual ZIP parsing for better compatibility
"""

import zipfile
import os
import zlib
import re
from typing import Dict, List, Any, Optional

class ImprovedDocxSkill:
    """An improved skill for reading and writing Word documents with better compatibility"""

    def __init__(self):
        self.doc_content = None
        self.files_data = {}

    def read_docx(self, file_path: str) -> Dict[str, Any]:
        """Read a .docx file and return its content with improved error handling"""
        try:
            if not os.path.exists(file_path):
                return {"error": f"File not found: {file_path}"}

            # Try standard zipfile first
            try:
                return self._read_with_zipfile(file_path)
            except zipfile.BadZipFile:
                # Fall back to manual parsing
                return self._read_with_manual_parsing(file_path)

        except Exception as e:
            return {"error": str(e)}

    def _read_with_zipfile(self, file_path: str) -> Dict[str, Any]:
        """Read using standard zipfile module"""
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            # Extract document.xml
            if 'word/document.xml' in zip_ref.namelist():
                with zip_ref.open('word/document.xml') as doc_file:
                    content = doc_file.read().decode('utf-8')
                    return self._parse_document_xml(content, file_path)
            else:
                return {"error": "Invalid docx file: missing document.xml"}

    def _read_with_manual_parsing(self, file_path: str) -> Dict[str, Any]:
        """Read using manual ZIP parsing for corrupted files"""
        try:
            with open(file_path, 'rb') as f:
                data = f.read()

            # Find all local file headers
            pos = 0
            files_found = []

            while pos < len(data) - 30:
                if data[pos:pos+4] == b'PK\x03\x04':
                    # Parse local file header
                    version = int.from_bytes(data[pos+4:pos+6], 'little')
                    flags = int.from_bytes(data[pos+6:pos+8], 'little')
                    compression = int.from_bytes(data[pos+8:pos+10], 'little')
                    compressed_size = int.from_bytes(data[pos+18:pos+22], 'little')
                    uncompressed_size = int.from_bytes(data[pos+22:pos+26], 'little')
                    filename_length = int.from_bytes(data[pos+26:pos+28], 'little')
                    extra_length = int.from_bytes(data[pos+28:pos+30], 'little')

                    filename = data[pos+30:pos+30+filename_length].decode('utf-8', errors='ignore')

                    files_found.append({
                        'name': filename,
                        'pos': pos,
                        'compression': compression,
                        'compressed_size': compressed_size,
                        'uncompressed_size': uncompressed_size,
                        'data_start': pos + 30 + filename_length + extra_length
                    })

                    # Move to next file
                    pos = pos + 30 + filename_length + extra_length + compressed_size
                else:
                    pos += 1

            # Find and extract document.xml
            for f in files_found:
                if 'word/document.xml' in f['name']:
                    if f['compression'] == 8:  # DEFLATE
                        compressed_data = data[f['data_start']:f['data_start']+f['compressed_size']]
                        decompressed = zlib.decompress(compressed_data, -15)
                        content = decompressed.decode('utf-8')
                        return self._parse_document_xml(content, file_path)
                    break

            return {"error": "Could not extract document content"}

        except Exception as e:
            return {"error": f"Manual parsing failed: {str(e)}"}

    def _parse_document_xml(self, xml_content: str, file_path: str) -> Dict[str, Any]:
        """Parse document.xml and extract text content"""
        try:
            # Extract text content using regex
            text_pattern = r'<w:t[^>]*>(.*?)</w:t>'
            texts = re.findall(text_pattern, xml_content)

            # Extract paragraph structure
            para_pattern = r'<w:p[^>]*>(.*?)</w:p>'
            paragraphs_raw = re.findall(para_pattern, xml_content, re.DOTALL)

            paragraphs = []
            for para in paragraphs_raw:
                # Extract text from paragraph
                para_texts = re.findall(text_pattern, para)
                if para_texts:
                    full_text = ''.join(para_texts)
                    if full_text.strip():
                        # Try to detect style
                        style_match = re.search(r'<w:pStyle w:val="([^"]*)"', para)
                        style = style_match.group(1) if style_match else "Normal"
                        paragraphs.append({
                            "text": full_text.strip(),
                            "style": style
                        })

            return {
                "success": True,
                "file_path": file_path,
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "table_count": 0,  # Manual parsing doesn't easily extract tables
                "raw_text": ' '.join(texts)
            }

        except Exception as e:
            return {"error": f"XML parsing failed: {str(e)}"}

    def write_docx(self, file_path: str, content: Any, title: Optional[str] = None) -> Dict[str, Any]:
        """Write content to a .docx file"""
        try:
            from docx import Document

            doc = Document()

            # Add title if provided
            if title:
                doc.add_heading(title, 0)

            # Add content
            if isinstance(content, str):
                # Split by paragraphs
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            elif isinstance(content, list):
                # List of paragraphs
                for para_text in content:
                    if isinstance(para_text, str) and para_text.strip():
                        doc.add_paragraph(para_text.strip())
                    elif isinstance(para_text, dict):
                        # Handle structured content
                        if 'heading' in para_text:
                            level = para_text.get('level', 1)
                            doc.add_heading(para_text['heading'], level=level)
                        elif 'text' in para_text:
                            doc.add_paragraph(para_text['text'])

            # Save the document
            doc.save(file_path)

            return {
                "success": True,
                "file_path": file_path,
                "message": f"Document saved successfully to {file_path}"
            }

        except Exception as e:
            return {"error": str(e)}

    def add_heading(self, text: str, level: int = 1) -> None:
        """Add a heading to the current document (for build mode)"""
        # This would require maintaining state for incremental document building
        # For now, use write_docx with structured content
        pass

    def add_paragraph(self, text: str, style: Optional[str] = None) -> None:
        """Add a paragraph to the current document (for build mode)"""
        # This would require maintaining state for incremental document building
        # For now, use write_docx with structured content
        pass

    def save(self, file_path: str) -> Dict[str, Any]:
        """Save the current document (for build mode)"""
        # This would require maintaining state for incremental document building
        # For now, use write_docx with structured content
        return {"error": "Use write_docx method instead"}

# Example usage
if __name__ == "__main__":
    skill = ImprovedDocxSkill()

    # Example: Read a docx file
    doc_path = "/sessions/nice-inspiring-mayer/mnt/CampusAgent/HKCampus_RAG面试备战手册.docx"
    result = skill.read_docx(doc_path)

    if "error" not in result:
        print(f"Successfully read: {result['file_path']}")
        print(f"Paragraphs found: {result['paragraph_count']}")
        print(f"Raw text length: {len(result.get('raw_text', ''))} characters")

        print("\nFirst 10 paragraphs:")
        for i, para in enumerate(result['paragraphs'][:10]):
            print(f"  {i+1}. [{para['style']}] {para['text'][:80]}...")
    else:
        print(f"Error: {result['error']}")

    # Example: Write a docx file
    content = [
        "This is a test document created by ImprovedDocxSkill.",
        "",
        {"heading": "Chapter 1: Introduction", "level": 1},
        "This is the introduction section of the document.",
        "",
        {"heading": "Chapter 2: Features", "level": 1},
        "The ImprovedDocxSkill supports:",
        "- Reading existing .docx files (including corrupted ones)",
        "- Writing new .docx files",
        "- Better error handling and fallback mechanisms"
    ]

    output_path = "/sessions/nice-inspiring-mayer/mnt/outputs/improved_test_output.docx"
    result = skill.write_docx(output_path, content, title="ImprovedDocxSkill Test")

    if "error" not in result:
        print(f"\nSuccessfully wrote: {result['file_path']}")
        print(f"Message: {result['message']}")
    else:
        print(f"\nError writing: {result['error']}")
